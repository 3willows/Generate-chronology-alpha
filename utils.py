import re
import csv
import pandas as pd
from docx import Document
import os
from fileinput import filename
import pypandoc 
from datetime import datetime
from flask import *  

app = Flask(__name__)   

def clean_markdown_document(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as file:
        content = file.read()

    # Find the position of the first occurrence of "1."
    first_occurrence_index = content.find("1.")
    # Find the position of the first occurrence of a footnote
    first_occurrence_of_footnote = content.find("[^1]:")  

    # Determine the index to cut off the content
    if first_occurrence_index!= -1:
        if first_occurrence_of_footnote!= -1 and first_occurrence_of_footnote > first_occurrence_index:
            cleaned_content = content[:first_occurrence_of_footnote]  # Slice content up to the footnote
        else:
            cleaned_content = content[first_occurrence_index:]  # Keep content from "1." onwards
    else:
        cleaned_content = content  # If "1." is not found, keep the content as is

    with open(output_file, 'w', encoding='utf-8') as file:
        file.write(cleaned_content)

    print(f"Markdown document cleaned and saved as {output_file}")


def extract_numbered_items_to_csv(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as file:
        content = file.read()

    # Regex pattern to find numbered items
    pattern = r'(\d+)\.\s*(.*?)\n(?=\d+\.\s|$)'  # Matches numbers followed by text, ending with a new line before another number

    # Find all matches
    matches = re.findall(pattern, content, re.DOTALL)

    # Write the matches to a CSV file
    with open(output_file, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Number', 'Text'])  # Writing header

        for number, text in matches:
            writer.writerow([number, text.strip()])  # Writing each row with number and text

    print(f"Data has been extracted and saved to {output_file}")


def parse_date(date_str):
    # Attempt to parse the date string into a datetime object
    date_formats = [
        '%m/%d/%Y', '%m-%d-%Y', '%m/%d/%y', '%m-%d-%y',
        '%d %B %Y', '%d %b %Y', '%B %d, %Y', '%b %d, %Y',
        '%Y-%m-%d'
    ]
    
    for date_format in date_formats:
        try:
            return datetime.strptime(date_str, date_format)
        except ValueError:
            pass
    
    return None

def extract_dates(input_file, output_file):
    extracted_data = []

    # Define regular expression patterns for common date formats
    date_patterns = [
        r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b',  # MM/DD/YYYY or MM-DD-YYYY or MM/DD/YY or MM-DD-YY
        r'\b(\d{1,2} [A-Za-z]{3,9} \d{2,4})\b',  # DD Month YYYY
        r'\b([A-Za-z]{3,9} \d{1,2}, \d{2,4})\b',  # Month DD, YYYY
        r'\b(\d{4}-\d{2}-\d{2})\b'  # YYYY-MM-DD
    ]

    with open(input_file, 'r', encoding='utf-8') as csvfile:
        reader = csv.DictReader(csvfile)
        
        for row in reader:
            paragraph_number = row['Number']  # Get the paragraph number
            text = row['Text']  # Get the text
            
            for pattern in date_patterns:
                matches = re.findall(pattern, text)
                for match in matches:
                    try:
                        # Attempt to parse the match into a datetime object
                        date = parse_date(match)
                        if date:
                            extracted_data.append({'Date': date, 'Text': text, 'Paragraph Number': paragraph_number})
                    except ValueError:
                        # Ignore if the match cannot be parsed into a valid date
                        pass

    extracted_data.sort(key=lambda x: x['Date'])

    with open(output_file, 'w', newline='', encoding='utf-8') as csvfile:
        fieldnames = ['Date', 'Text', 'Paragraph Number']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        
        writer.writeheader()
        
        for entry in extracted_data:
            writer.writerow({'Date': entry['Date'].strftime("%Y-%m-%d"), 
                             'Text': entry['Text'], 
                             'Paragraph Number': entry['Paragraph Number']})

    print(f"Date extraction completed and saved to {output_file}")

def create_word_document_from_csv(input_file, output_file):
    data = pd.read_csv(input_file)

    # Create a new Word document
    doc = Document()

    # Add a title to the document (optional)
    doc.add_heading('Draft Chronology', level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'  # Optional: set a table style

    # Add the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Text'
    hdr_cells[2].text = 'Paragraph Number'

    for index, row in data.iterrows():
        row_cells = table.add_row().cells  # Add a new row
        row_cells[0].text = str(row['Date'])  # Fill Date
        row_cells[1].text = str(row['Text'])  # Fill Text
        row_cells[2].text = str(row['Paragraph Number'])  # Fill Paragraph Number

    doc.save(output_file)

    print(f"Word document '{output_file}' created successfully.")

def everything_function(f):
    try:
        pypandoc.convert_file(f, 'md', outputfile='input.md')
        clean_markdown_document('input.md', 'cleaned.md')
        extract_numbered_items_to_csv('cleaned.md', 'all_dates.csv')
        extract_dates('all_dates.csv', 'dates_extracted.csv')
        create_word_document_from_csv('dates_extracted.csv', 'draft-chronology.docx')

    finally:
        # Remove intermediate files
        for file in [f, 'input.md', 'cleaned.md', 'all_dates.csv', 'dates_extracted.csv']:
            if os.path.exists(file):
                os.remove(file)
                print(f"Removed {file}")

@app.route('/')   
def main():
    everything_function() 

if __name__ == '__main__':   
    app.run(debug=True)